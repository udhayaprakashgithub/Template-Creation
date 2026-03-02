from pathlib import Path

from django.core.files.base import File
from docx import Document

from ..models import GeneratedDocument


class DocumentBuilder:
    def populate_template(self, template_file_path, replacements: dict, output_path):
        doc = Document(template_file_path)
        self._replace_in_paragraphs(doc.paragraphs, replacements)
        self._replace_in_tables(doc, replacements)

        for section in doc.sections:
            self._replace_in_paragraphs(section.header.paragraphs, replacements)
            self._replace_in_paragraphs(section.footer.paragraphs, replacements)

        doc.save(output_path)
        return output_path

    def create_generated_record(self, source_document, template, user, output_path, is_final=False):
        with Path(output_path).open("rb") as file_handle:
            generated = GeneratedDocument.objects.create(
                document=source_document,
                template=template,
                generated_by=user,
                is_final=is_final,
            )
            generated.generated_file.save(Path(output_path).name, File(file_handle), save=True)
            return generated

    @staticmethod
    def _replace_in_paragraphs(paragraphs, replacements):
        for paragraph in paragraphs:
            for key, value in replacements.items():
                placeholder = f"<{key}>"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _replace_in_tables(self, document, replacements):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_paragraphs(cell.paragraphs, replacements)
